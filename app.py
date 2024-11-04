import streamlit as st
import requests
import io
import fitz  # PyMuPDF
import re
import os
from docx import Document
from docx.shared import Inches

# OCR.space API endpoint
OCR_SPACE_API_URL = 'https://api.ocr.space/parse/image'
# Get your API key from the environment variable
OCR_SPACE_API_KEY = os.getenv('OCR_SPACE_API_KEY')
if not OCR_SPACE_API_KEY:
    st.error("OCR_SPACE_API_KEY environment variable not set.")
    st.stop()

# URL of the PDF document
pdf_url = "https://resource.cdn.icai.org/69249asb55316-as21.pdf"

def download_pdf(url):
    response = requests.get(url)
    response.raise_for_status()
    return io.BytesIO(response.content)

def pdf_to_images(pdf_file):
    doc = fitz.open(stream=pdf_file.read(), filetype="pdf")
    images = []
    for page_num in range(len(doc)):
        page = doc.load_page(page_num)
        pix = page.get_pixmap()
        img_data = pix.tobytes("png")
        images.append(img_data)
    return images

def ocr_image_with_api(image_data):
    payload = {
        'isOverlayRequired': False,
        'apikey': OCR_SPACE_API_KEY,
        'language': 'eng',
    }
    files = {
        'filename': ('image.png', image_data),
    }
    response = requests.post(OCR_SPACE_API_URL, data=payload, files=files)
    result = response.json()
    if result.get('IsErroredOnProcessing'):
        st.error("Error during OCR processing.")
        return ''
    return result['ParsedResults'][0]['ParsedText']

def ocr_images(images):
    text = ""
    for idx, image_data in enumerate(images):
        st.write(f"Performing OCR on page {idx+1}...")
        page_text = ocr_image_with_api(image_data)
        text += page_text + "\n"
    return text

def parse_steps(text):
    """
    Parses the OCR-extracted text to find steps and their content.
    Adjust the regex pattern based on the document's format.
    """
    # Regular expression to match steps starting with "Step X:"
    pattern = r"(?:^|\n)(Step\s+\d+[\.:]?\s*)(.*?)(?=\nStep\s+\d+[\.:]?|$)"
    matches = re.findall(pattern, text, re.DOTALL | re.IGNORECASE)

    steps = []
    for match in matches:
        step_title = match[0].strip()
        step_content = match[1].strip()
        steps.append((step_title, step_content))
    return steps

def create_word_document(steps):
    """
    Creates a Word document with the extracted steps.
    Returns the BytesIO object containing the document data.
    """
    doc = Document()
    doc.add_heading('Extracted Steps', 0)

    for step_title, step_content in steps:
        doc.add_heading(step_title, level=1)
        doc.add_paragraph(step_content)

    # Save the document to a BytesIO object
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io

def main():
    st.title("Process PDF Steps from URL using OCR.space API")

    st.write("Downloading the PDF document...")
    try:
        pdf_file = download_pdf(pdf_url)
    except requests.exceptions.RequestException as e:
        st.error(f"Error downloading PDF: {e}")
        return

    st.write("Converting PDF pages to images...")
    images = pdf_to_images(pdf_file)

    if not images:
        st.error("Failed to convert PDF pages to images.")
        return

    st.write("Performing OCR on the images to extract text...")
    ocr_text = ocr_images(images)

    if not ocr_text.strip():
        st.error("OCR did not extract any text from the images.")
        return

    st.write("Extracting steps from the OCR text...")
    steps = parse_steps(ocr_text)

    if not steps:
        st.error("No steps found in the document.")
        return

    st.success(f"Found {len(steps)} steps in the document.")

    # Create the Word document
    doc_file = create_word_document(steps)

    # Provide a download button for the Word document
    st.write("---")
    st.write("## Download Extracted Steps as Word Document")
    st.download_button(
        label="Download Word Document",
        data=doc_file,
        file_name="Extracted_Steps.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    # Optionally, display the steps in the app
    st.write("---")
    st.write("## Extracted Steps:")
    for i, (step_title, step_content) in enumerate(steps, 1):
        st.write(f"### {step_title}")
        st.write(step_content)
        st.write("---")

    # Optional: Execute steps
    # st.write("## Execute Steps:")
    # for i, (step_title, step_content) in enumerate(steps, 1):
    #     if st.button(f"Run {step_title}"):
    #         step_function = create_step_function(step_title, step_content, i)
    #         step_function()

if __name__ == "__main__":
    main()
